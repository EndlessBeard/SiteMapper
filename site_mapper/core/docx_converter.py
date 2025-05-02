import os
import glob
import json
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import logging
import uuid
import re
import urllib.parse
from django.conf import settings

logger = logging.getLogger(__name__)
color_fill = "FFC000"
def json_to_docx(json_path, docx_path=None):
    """
    Convert a site map JSON file to a DOCX document with a structured outline
    that can be collapsed/expanded in Word's outline view.
    
    Args:
        json_path (str): Path to the input JSON file
        docx_path (str): Path to save the output DOCX file (optional)
    """
    # If docx_path is not provided, create one based on the json_path
    if not docx_path:
        base_path = os.path.splitext(json_path)[0]
        docx_path = f"{base_path}.docx"
    
    # Load JSON data
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            site_data = json.load(f)
    except Exception as e:
        logger.error(f"Error loading JSON file {json_path}: {e}")
        site_data = []
    
    # Create new document
    doc = Document()
    
    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(0.1)
    section.right_margin = Inches(0.1)
    
    # Set default font and line spacing
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)  # Set font size to 10pt
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = Pt(11.0)  # Set line spacing to 0.75
    
    # Add title
    title = doc.add_heading('SITE MAP', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add document metadata
    doc.add_paragraph(f"Source: {json_path}")
    
    # Add content (example structure)
    if isinstance(site_data, list):
        for item in site_data:
            add_link_paragraph(doc, item, level=1)
    elif isinstance(site_data, dict):
        for key, value in site_data.items():
            doc.add_heading(key, level=1)
            doc.add_paragraph(str(value))
    
    # Save the document
    doc.save(docx_path)
    return docx_path


def process_link(doc, link, links_by_parent, level):
    """
    Process a link and its children recursively, adding them to the document.
    
    Args:
        doc: The document
        link: The current link to process
        links_by_parent: Dictionary of links indexed by parent UUID
        level: Current heading level
    """
    if not isinstance(link, dict) or 'url' not in link:
        return
    
    # Add this link as a heading
    add_link_heading(doc, link, level)
    
    # Process children if any
    if link.get('uuid') in links_by_parent:
        children = links_by_parent[link['uuid']]
        for child in children:
            process_link(doc, child, links_by_parent, level + 1)


def add_link_heading(doc, link, level):
    """
    Add a link to the document as a heading with appropriate formatting.
    
    Args:
        doc: The document
        link: The link data
        level: Heading level
    """
    # Don't go deeper than level 9 (Word's maximum heading level)
    if level > 9:
        level = 9
        
    url = link.get('url', '')
    link_text = link.get('link_text', url) or url
    
    # Determine link type and prefix
    link_type = link.get('type', 'page')
    
    # Create heading for this link
    heading = doc.add_heading('', level=level)
    
    # Add appropriate prefix based on link type with colored background
    prefix = ""
    if link_type == 'page':
        prefix = "[PAGE] "
        color_fill = "92D050"  # Green
    elif link_type == 'external':
        prefix = "[EXTERNAL] "
        color_fill = "00B0F0"  # Blue
    elif link_type in ['pdf', 'docx', 'xlsx', 'xls']:
        prefix = f"[{link_type.upper()}] "
        color_fill = "FFC000"  # Orange
    elif link_type == 'other':
        prefix = "[OTHER] "
        color_fill = "CCCCCC"  # Gray for unknown types
    elif link_type == 'broken':
        prefix = "[BROKEN] "
        color_fill = "FF0000"  # Red for broken links
    else:
        prefix = f"[{link_type.upper()}] "
        color_fill = "CCCCCC"  # Gray for unknown types
    
    # Create a run for the prefix with colored background
    run = heading.add_run(prefix)
    
    # Add the hyperlink
    add_hyperlink(heading, link_text, url)
    
    # Add depth information if available
    if 'depth' in link:
        heading.add_run(f" (Depth: {link['depth']})")


def process_page(doc, page, pages_by_parent, level):
    """
    Process a page and its children recursively, adding them to the document.
    
    Args:
        doc: The document
        page: Current page data
        pages_by_parent: Dictionary of pages indexed by parent IDs
        level: Current heading level
    """
    if not isinstance(page, dict):
        return
    
    add_link_paragraph(doc, page, level)
    
    # Process children if any
    page_id = page.get('id')
    if page_id and page_id in pages_by_parent:
        children = pages_by_parent[page_id]
        for child in children:
            process_page(doc, child, pages_by_parent, level + 1)


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        text: The text of the hyperlink
        url: The URL of the hyperlink
    """
    if not url:
        paragraph.add_run(text)
        return
        
    # This gets access to the document.xml.rels file and gets a new relation id
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the w:hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a w:r element
    new_run = OxmlElement('w:r')
    
    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')
    
    # Add color and underline to the w:rPr element
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # Add underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    # Join all the xml elements
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink


def add_toc(doc):
    """
    Add a table of contents to the document.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Table of contents with 3 levels, hyperlinks, and no page numbers
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(OxmlElement('w:t'))  # This is a placeholder for the TOC content
    r_element.append(fldChar3)


def process_job_to_docx(job_id):
    """
    Convert all site map JSON files for a job into DOCX format.
    
    Args:
        job_id: The ID of the SiteMapJob
        
    Returns:
        list: Paths to the created DOCX files
    """
    logger.info(f"Converting site maps to DOCX for job {job_id}")
    
    # Find all site map JSON files for this job
    job_dir = os.path.join(settings.MEDIA_ROOT, 'site_mapper', f'job_{job_id}')
    if not os.path.exists(job_dir):
        logger.error(f"Job directory not found: {job_dir}")
        return []
    
    # Find all site map JSON files
    json_files = glob.glob(os.path.join(job_dir, 'site_map_*.json'))
    logger.info(f"Found {len(json_files)} site map JSON files in {job_dir}")
    
    if not json_files:
        logger.warning(f"No site map JSON files found for job {job_id}")
        return []
    
    # Convert each JSON file to DOCX using the correct converter function
    docx_files = []
    for json_file in json_files:
        docx_file = json_file.replace('.json', '.docx')
        try:
            logger.info(f"Converting {json_file} to {docx_file}")
            # Use convert_json_to_docx instead of json_to_docx
            filename = os.path.basename(json_file)
            start_url = extract_url_from_filename(filename)
            convert_json_to_docx(json_file, docx_file, start_url)
            docx_files.append(docx_file)
            logger.info(f"Successfully created {docx_file}")
        except Exception as e:
            logger.error(f"Error converting {json_file} to DOCX: {e}", exc_info=True)
    
    return docx_files


def extract_url_from_filename(filename):
    """
    Extract the original URL from a sanitized filename.
    This function handles file names like 'site_map_www_research_uky_edu_collaborative-grant-services_about-us.json'
    
    Args:
        filename (str): Sanitized filename like 'site_map_www_example_com.json'
        
    Returns:
        str: Best guess at the original URL
    """
    # Remove 'site_map_' prefix and file extension
    if filename.startswith('site_map_'):
        url_part = filename[9:]  # Remove 'site_map_'
    else:
        url_part = filename
        
    # Remove file extension if present
    url_part = os.path.splitext(url_part)[0]
    
    # Replace underscores with dots only for domain part
    parts = url_part.split('_')
    if len(parts) >= 2:
        # First few parts are likely domain components (www.example.com)
        domain_parts = []
        path_parts = []
        
        # Assume first part and second part belong to domain
        domain_parts = parts[:2]
        
        # Check if third part might still be domain (like 'edu', 'org', etc.)
        if len(parts) > 2 and len(parts[2]) <= 3:  # Common TLDs are short
            domain_parts.append(parts[2])
            path_parts = parts[3:]
        else:
            path_parts = parts[2:]
        
        # Join domain parts with dots
        domain = '.'.join(domain_parts)
        
        # Join path parts with hyphens/slashes based on patterns
        path = '/'.join([p.replace('-', '/') if '/' in p else p for p in path_parts])
        
        # If path doesn't start with slash and isn't empty, add one
        if path and not path.startswith('/'):
            path = '/' + path
            
        url_part = domain + path
    else:
        # Simple replacement if we can't identify parts
        url_part = url_part.replace('_', '.')
    
    # If it doesn't start with http or https, add it
    if not url_part.startswith(('http://', 'https://')):
        url_part = f"https://{url_part}"
    
    return url_part


def convert_json_to_docx(json_path, docx_path, start_url=None):
    """
    Convert a site map JSON file to a DOCX document with a structured outline.
    
    Args:
        json_path (str): Path to the input JSON file
        docx_path (str): Path to save the output DOCX file
        start_url (str, optional): The starting URL for this site map
        
    Returns:
        str: Path to the created DOCX file
    """
    # Load JSON data
    with open(json_path, 'r', encoding='utf-8') as f:
        site_data = json.load(f)
    
    # Create new document
    doc = Document()
    
    # Get the domain from the start_url for the title
    domain = "SITE MAP"
    if start_url:
        try:
            parsed_url = urllib.parse.urlparse(start_url)
            domain = parsed_url.netloc
            path = parsed_url.path.rstrip('/')
            if path:
                domain += path
        except:
            pass
    
    # Add title with the starting URL domain
    title = doc.add_heading(f'SITE MAP: {domain}', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add document metadata
    doc.add_paragraph(f"Starting URL: {start_url or 'N/A'}")
    
    # Continue with the rest of the document creation
    # (using the existing code from json_to_docx function)
    
    # Add color legend
    doc.add_heading('Color Legend', level=1)
    
    # Create a table for the legend with colored backgrounds
    legend_table = doc.add_table(rows=4, cols=1)
    legend_table.style = 'Table Grid'
    
    # Internal links - green
    cell = legend_table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("[PAGE] Internal Page Links")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="92D050"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # External links - blue
    cell = legend_table.rows[1].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("[EXTERNAL] External Links")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="00B0F0"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Resource links - orange
    cell = legend_table.rows[2].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("[DOC] Resource Links (PDFs, DOCXs, etc.)")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFC000"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Broken links - red
    cell = legend_table.rows[3].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("[BROKEN] Broken Links")
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FF0000"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    # Process site structure using existing functions
    try:
        # Continue with the existing structure processing code
        if isinstance(site_data, dict) and 'roots' in site_data:
            # Handle the new format with "roots" key
            doc.add_heading('Site Structure', level=1)
            for root in site_data.get('roots', []):
                process_root_node(doc, root, 2)
        elif isinstance(site_data, list):
            # Handle the old format (list of links)
            # This code is from the original json_to_docx function
            # ... (Copy the relevant code here)
            pass
        else:
            # Handle other formats as in the original function
            # ... (Copy the relevant code here)
            pass
    except Exception as e:
        logging.error(f"Error processing site structure: {str(e)}", exc_info=True)
        doc.add_paragraph(f"Error processing site structure: {str(e)}")
    
    # Save the document
    doc.save(docx_path)
    
    return docx_path


def process_root_node(doc, root, level):
    """
    Process a root node and its children recursively.
    
    Args:
        doc: The document
        root: The root node dictionary
        level: Current heading level
    """
    # Add the root URL as a heading
    add_link_heading(doc, root, level)
    
    # Process children
    if 'children' in root and root['children']:
        for child in root['children']:
            process_root_node(doc, child, level + 1)


def process_all_job_files(job_id):
    """
    Convert all JSON site maps in a job to DOCX files.
    
    Args:
        job_id (str or int): ID of the job
        
    Returns:
        list: Paths to the created DOCX files
    """
    job_dir = os.path.join('media', f'site_mapper/job_{job_id}')
    if not os.path.exists(job_dir):
        logging.error(f"Job directory not found: {job_dir}")
        return []
    
    # Find all site map JSON files
    json_files = glob.glob(os.path.join(job_dir, 'site_map_*.json'))
    logging.info(f"Found {len(json_files)} site map JSON files for job {job_id}")
    
    # Debug output of found files
    for jf in json_files:
        logging.debug(f"Found JSON file: {os.path.basename(jf)}")
    
    # Process each JSON file
    docx_files = []
    for json_path in json_files:
        try:
            # Extract the starting URL from the filename
            filename = os.path.basename(json_path)
            start_url = extract_url_from_filename(filename)
            logging.info(f"Extracted URL from filename: {start_url}")
            
            # Create DOCX with the same naming convention
            docx_path = os.path.splitext(json_path)[0] + '.docx'
            
            # Convert JSON to DOCX
            result_path = convert_json_to_docx(json_path, docx_path, start_url)
            
            docx_files.append(result_path)
            logging.info(f"Created DOCX file: {result_path}")
        except Exception as e:
            logging.error(f"Error processing {json_path}: {str(e)}", exc_info=True)
    
    return docx_files


def check_and_fix_json(json_path):
    """
    Check if the JSON file is valid, and attempt to repair if needed.
    
    Args:
        json_path (str): Path to the JSON file
        
    Returns:
        dict or list: The loaded and potentially fixed JSON data
    """
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except json.JSONDecodeError as e:
        # Try to repair common JSON issues
        logging.warning(f"JSON decode error in {json_path}: {str(e)}")
        
        try:
            # Read the file as text
            with open(json_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Common fixes
            if content.strip() == "":
                # Empty file
                logging.warning(f"Empty JSON file: {json_path}")
                return {"links": [], "error": "Empty file"}
                
            # Try some basic repairs and reload
            fixed_content = content.replace("'", '"')  # Replace single quotes with double quotes
            fixed_content = re.sub(r',\s*}', '}', fixed_content)  # Remove trailing commas
            fixed_content = re.sub(r',\s*]', ']', fixed_content)
            
            data = json.loads(fixed_content)
            
            # Write the fixed content back to the file
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(data, f)
                
            return data
        except Exception as repair_error:
            logging.error(f"Failed to repair JSON: {str(repair_error)}")
            return {"links": [], "error": f"Invalid JSON: {str(e)}"}


def create_site_structure(links):
    """
    Create a hierarchical site structure from a flat list of links.
    
    Args:
        links (list): List of link dictionaries
        
    Returns:
        dict: Site structure with roots and hierarchical relationships
    """
    # Create a dictionary to hold link data by URL
    links_by_url = {}
    links_by_uuid = {}
    
    # First pass: index links by URL and UUID
    for link in links:
        if not isinstance(link, dict):
            continue
            
        url = link.get('url')
        if url:
            links_by_url[url] = link
            
        uuid_val = link.get('uuid')
        if uuid_val:
            links_by_uuid[uuid_val] = link
    
    # Second pass: build parent-child relationships
    roots = []
    for link in links:
        if not isinstance(link, dict):
            continue
            
        # Check if this is a root link (no parent or depth 0)
        is_root = False
        if 'parent' not in link or link['parent'] is None:
            is_root = True
        elif link.get('depth', 0) == 0:
            is_root = True
            
        if is_root:
            # Deep copy to avoid modifying the original
            root_link = link.copy()
            root_link['children'] = []
            roots.append(root_link)
            
    # Now add children to their parents
    for link in links:
        if not isinstance(link, dict):
            continue
            
        parent_id = link.get('parent')
        if parent_id and parent_id in links_by_uuid:
            # Find this parent in our tree
            for root in roots:
                parent_in_tree = find_node_by_uuid(root, parent_id)
                if parent_in_tree:
                    if 'children' not in parent_in_tree:
                        parent_in_tree['children'] = []
                    # Deep copy to avoid modifying the original
                    child_link = link.copy()
                    child_link['children'] = []
                    parent_in_tree['children'].append(child_link)
                    break
    
    return {"roots": roots, "links_by_uuid": links_by_uuid}


def find_node_by_uuid(node, uuid_val):
    """
    Find a node in the tree by UUID.
    
    Args:
        node (dict): The current node to check
        uuid_val (str): UUID to look for
        
    Returns:
        dict or None: The found node or None
    """
    if node.get('uuid') == uuid_val:
        return node
        
    if 'children' in node:
        for child in node['children']:
            found = find_node_by_uuid(child, uuid_val)
            if found:
                return found
                
    return None


def extract_links_from_data(data, links=None):
    """
    Recursively extract links from any data structure.
    
    Args:
        data: Any data structure that might contain links
        links: List to accumulate links (used in recursion)
        
    Returns:
        list: Extracted links
    """
    if links is None:
        links = []
    
    if isinstance(data, dict):
        if 'url' in data:
            links.append(data)
        for value in data.values():
            extract_links_from_data(value, links)
    elif isinstance(data, list):
        for item in data:
            extract_links_from_data(item, links)
            
    return links


def add_link_paragraph(doc, link, level=2):
    """
    Add a link to the document as a paragraph with appropriate formatting.
    
    Args:
        doc: The document
        link: The link data
        level: Heading level (default 2)
    """
    if not isinstance(link, dict) or 'url' not in link:
        return
        
    url = link.get('url', '')
    link_text = link.get('link_text', url) or url
    
    # Determine link type and prefix
    link_type = link.get('type', 'page')
    
    # Create paragraph for this link
    paragraph = doc.add_paragraph()
    paragraph.style = f'Heading {level}'
    
    # Add indentation based on depth (4 spaces per level)
    depth = link.get('depth', 0)
    indent = ' ' * (4 * depth)  # 4 spaces per depth level
    
    # Add appropriate prefix based on link type
    prefix = ""
    if link_type == 'page':
        prefix = f"{indent}[PAGE] "
        color_fill = "92D050"  # Green
    elif link_type == 'external':
        prefix = f"{indent}[EXTERNAL] "
        color_fill = "00B0F0"  # Blue
    elif link_type in ['pdf', 'docx', 'xlsx', 'xls']:
        prefix = f"{indent}[{link_type.upper()}] "
        color_fill = "FFC000"  # Orange
    elif link_type == 'broken':
        prefix = f"{indent}[BROKEN] "
        color_fill = "FF0000"  # Red
    else:
        prefix = f"{indent}[{link_type.upper()}] "
        color_fill = "CCCCCC"  # Gray for unknown types
    
    # Add prefix
    paragraph.add_run(prefix)
    
    # Add the hyperlink
    add_hyperlink(paragraph, link_text, url)
    
    # Add depth information if available
    if 'depth' in link:
        paragraph.add_run(f" (Depth: {link['depth']})")
    
    # Add color fill to the paragraph background
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_fill}"/>')
    paragraph._p.get_or_add_pPr().append(shading_elm)