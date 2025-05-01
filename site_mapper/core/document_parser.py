"""
Document Parser module for Site_Mapper2.
This module extracts text and links from various document types.
"""
import logging
import re
import os
from urllib.parse import urljoin
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text
import docx
import openpyxl

logger = logging.getLogger(__name__)

# URL pattern for matching various URL formats
URL_PATTERN = re.compile(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+(?:/[-\w./?%&=]*)?')

class DocumentParser:
    """
    Extracts text and links from various document types.
    Supports PDF, Word, and Excel formats.
    """
    
    def __init__(self, base_url=None):
        """
        Initialize the DocumentParser.
        
        Args:
            base_url (str, optional): Base URL for resolving relative links
        """
        self.base_url = base_url
    
    def parse(self, file_path):
        """
        Parse a document file and extract text and links.
        
        Args:
            file_path (str): Path to the document
            
        Returns:
            dict: Contains extracted text and links
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {'text': '', 'links': []}
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._parse_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._parse_xlsx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                return {'text': '', 'links': []}
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return {'text': '', 'links': []}
    
    def _parse_pdf(self, file_path):
        """
        Parse a PDF document and extract text and links.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            dict: Contains extracted text and links
        """
        logger.info(f"Parsing PDF: {file_path}")
        
        # Try PyPDF2 first
        extracted_text = ""
        urls = []
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    extracted_text += page_text + "\n\n"
                    
                    # Extract links from annotations if available
                    if '/Annots' in page:
                        annotations = page['/Annots']
                        for annotation in annotations:
                            if annotation.get('/Subtype') == '/Link' and '/A' in annotation:
                                if '/URI' in annotation['/A']:
                                    urls.append({
                                        'url': annotation['/A']['/URI'],
                                        'text': '',
                                        'type': 'pdf_link'
                                    })
            
            # If PyPDF2 didn't extract enough text, try pdfminer
            if len(extracted_text.strip()) < 100:
                extracted_text = pdfminer_extract_text(file_path)
                
        except Exception as e:
            logger.warning(f"Error with PyPDF2, falling back to pdfminer: {e}")
            try:
                extracted_text = pdfminer_extract_text(file_path)
            except Exception as e2:
                logger.error(f"Failed to extract text from PDF: {e2}")
                return {'text': '', 'links': []}
                
        # Find URLs in the text
        text_urls = URL_PATTERN.findall(extracted_text)
        for url in text_urls:
            if url not in [u['url'] for u in urls]:
                urls.append({
                    'url': url,
                    'text': '',
                    'type': 'pdf_text_link'
                })
                
        return {'text': extracted_text, 'links': urls}
    
    def _parse_docx(self, file_path):
        """
        Parse a Word document and extract text and links.
        
        Args:
            file_path (str): Path to the Word file
            
        Returns:
            dict: Contains extracted text and links
        """
        logger.info(f"Parsing Word document: {file_path}")
        
        extracted_text = ""
        urls = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
                
                # Extract hyperlinks from paragraph
                for run in para.runs:
                    if run._element.xpath('.//w:hyperlink'):
                        for hyperlink in run._element.xpath('.//w:hyperlink'):
                            for rel_id in hyperlink.values():
                                if rel_id.startswith('rId'):
                                    try:
                                        target = doc.part.rels[rel_id].target_ref
                                        urls.append({
                                            'url': target,
                                            'text': run.text,
                                            'type': 'docx_link'
                                        })
                                    except KeyError:
                                        pass
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        extracted_text += cell_text + "\n"
                        
            # Find URLs in the text
            text_urls = URL_PATTERN.findall(extracted_text)
            for url in text_urls:
                if url not in [u['url'] for u in urls]:
                    urls.append({
                        'url': url,
                        'text': '',
                        'type': 'docx_text_link'
                    })
                    
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            return {'text': '', 'links': []}
            
        return {'text': extracted_text, 'links': urls}
    
    def _parse_xlsx(self, file_path):
        """
        Parse an Excel document and extract text and links.
        
        Args:
            file_path (str): Path to the Excel file
            
        Returns:
            dict: Contains extracted text and links
        """
        logger.info(f"Parsing Excel document: {file_path}")
        
        extracted_text = ""
        urls = []
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                extracted_text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows():
                    row_text = ""
                    for cell in row:
                        if cell.value:
                            cell_value = str(cell.value)
                            row_text += cell_value + "\t"
                            
                            # Check if cell contains a hyperlink
                            if cell.hyperlink:
                                url = cell.hyperlink.target
                                if url:
                                    urls.append({
                                        'url': url,
                                        'text': cell_value,
                                        'type': 'xlsx_link'
                                    })
                                    
                    extracted_text += row_text.strip() + "\n"
                extracted_text += "\n"
                
            # Find URLs in the text
            text_urls = URL_PATTERN.findall(extracted_text)
            for url in text_urls:
                if url not in [u['url'] for u in urls]:
                    urls.append({
                        'url': url,
                        'text': '',
                        'type': 'xlsx_text_link'
                    })
                    
        except Exception as e:
            logger.error(f"Failed to extract text from Excel document: {e}")
            return {'text': '', 'links': []}
            
        return {'text': extracted_text, 'links': urls}